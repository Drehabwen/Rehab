from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime
from typing import Dict


class DocumentGenerator:
    def __init__(self, config: Dict):
        self.config = config
        self.exports_dir = config["exports_dir"]
        self._ensure_directories()

    def _ensure_directories(self):
        if not os.path.exists(self.exports_dir):
            os.makedirs(self.exports_dir)

    def generate_word(self, case: Dict) -> str:
        doc = Document()
        
        self._setup_document_styles(doc)
        self._add_header(doc)
        self._add_patient_info(doc, case)
        self._add_case_content(doc, case)
        self._add_footer(doc)
        
        filename = f"{case['case_id']}_{case['patient_name']}_病历.docx"
        filepath = os.path.join(self.exports_dir, filename)
        doc.save(filepath)
        
        return filepath

    def _setup_document_styles(self, doc):
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(12)

    def _add_header(self, doc):
        hospital_name = self.config.get("hospital_name", "医疗中心")
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run(hospital_name)
        header_run.font.size = Pt(18)
        header_run.font.bold = True
        header_run.font.name = '宋体'
        header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        doc.add_paragraph()
        doc.add_paragraph()

    def _add_patient_info(self, doc, case):
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        
        labels = ["姓名", "性别", "年龄", "就诊日期"]
        values = [
            case.get("patient_name", ""),
            case.get("gender", ""),
            str(case.get("age", "")) + "岁",
            case.get("visit_date", "")
        ]
        
        for i in range(4):
            cell = table.rows[0].cells[i]
            cell.text = labels[i]
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cell = table.rows[1].cells[i]
            cell.text = values[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()

    def _add_case_content(self, doc, case):
        sections = [
            ("主诉", case.get("chief_complaint", "")),
            ("现病史", case.get("present_illness", "")),
            ("既往史", case.get("past_history", "")),
            ("过敏史", case.get("allergies", "")),
            ("体格检查", case.get("physical_exam", "")),
            ("诊断", case.get("diagnosis", "")),
            ("治疗建议", case.get("treatment_plan", ""))
        ]
        
        for title, content in sections:
            if content:
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(title + "：")
                title_run.font.bold = True
                title_run.font.size = Pt(12)
                
                content_para = doc.add_paragraph(content)
                content_para.paragraph_format.left_indent = Inches(0.3)
                
                doc.add_paragraph()

    def _add_footer(self, doc):
        doctor_name = self.config.get("doctor_name", "医生")
        current_time = datetime.now().strftime("%Y年%m月%d日 %H:%M")
        
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_para.add_run(f"医生签名：{doctor_name}")
        footer_para.add_run("\n")
        footer_para.add_run(f"记录时间：{current_time}")
        
        doc.add_paragraph()
