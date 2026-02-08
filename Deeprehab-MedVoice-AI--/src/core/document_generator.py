from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import re
from datetime import datetime
from typing import Dict
try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False


class DocumentGenerator:
    def __init__(self, config: Dict):
        self.config = config
        self.exports_dir = config["exports_dir"]
        self._ensure_directories()

    def _clean_markdown(self, text: str) -> str:
        """清理 Markdown 符号，返回纯文本"""
        if not text:
            return ""
        
        # 0. 移除常见的 Markdown 特殊块
        text = re.sub(r'```[\s\S]*?```', '', text)  # 代码块
        text = re.sub(r'`(.*?)`', r'\1', text)      # 行内代码
        
        # 1. 移除标题符号 (如 # 标题)
        text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
        
        # 2. 移除加粗和斜体 (***bold-italic***, **bold**, *italic*)
        # 注意顺序：先长后短
        text = re.sub(r'(\*\*\*|___)(.*?)\1', r'\2', text)
        text = re.sub(r'(\*\*|__)(.*?)\1', r'\2', text)
        text = re.sub(r'(\*|_)(.*?)\1', r'\2', text)
        
        # 3. 移除列表符号 (如 - item, * item, + item, 1. item)
        # 增加对没有空格的列表符的兼容，以及对特殊星号的清理
        text = re.sub(r'^\s*[-*+•✳]\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*\d+\.\s*', '', text, flags=re.MULTILINE)
        
        # 4. 移除水平线 (---, ***, ___)
        text = re.sub(r'^\s*[-*_]{3,}\s*$', '', text, flags=re.MULTILINE)
        
        # 5. 移除引用符号 (> quote)
        text = re.sub(r'^>\s+', '', text, flags=re.MULTILINE)
        
        # 6. 移除链接 [text](url) -> text
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        
        # 7. 最后清理多余的空格和换行
        lines = [line.strip() for line in text.split('\n')]
        # 过滤掉连续的空行，只保留必要的换行
        result_lines = []
        for i, line in enumerate(lines):
            if line or (i > 0 and lines[i-1]):
                result_lines.append(line)
        
        return '\n'.join(result_lines).strip()

    def _ensure_directories(self):
        if not os.path.exists(self.exports_dir):
            os.makedirs(self.exports_dir)

    def generate_word(self, case: Dict) -> str:
        doc = Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        self._setup_document_styles(doc)
        self._add_header(doc)
        self._add_patient_info(doc, case)
        self._add_divider(doc)
        self._add_case_content(doc, case)
        self._add_footer(doc, case)
        
        filename = f"{case['case_id']}_{case['patient_name']}_病历.docx"
        filepath = os.path.join(self.exports_dir, filename)
        doc.save(filepath)
        
        return filepath

    def generate_pdf(self, case: Dict) -> str:
        if not FPDF_AVAILABLE:
            raise ImportError("fpdf2 is not installed. Please install it to use PDF export.")
            
        filename = f"{case['case_id']}_{case['patient_name']}_病历.pdf"
        filepath = os.path.join(self.exports_dir, filename)
        
        pdf = FPDF()
        pdf.add_page()
        
        # 注册并设置中文字体
        font_path_hei = r"C:\Windows\Fonts\simhei.ttf"
        
        # 检查字体文件是否存在，如果不存在则尝试当前目录下
        if not os.path.exists(font_path_hei):
            font_path_hei = "simhei.ttf"
            
        pdf.add_font("SimHei", "", font_path_hei)
        pdf.add_font("SimSun", "", font_path_hei) # 暂时使用黑体代替宋体以避开 TTC 兼容性问题
        
        # 医院名称
        hospital_name = self.config.get("hospital_name", "XX社区卫生服务中心")
        pdf.set_font("SimHei", size=18)
        pdf.cell(0, 15, hospital_name, ln=True, align="C")
        
        # 标题
        pdf.set_font("SimHei", size=16)
        pdf.cell(0, 10, "门诊病历", ln=True, align="C")
        pdf.ln(5)
        
        # 患者信息表格
        pdf.set_font("SimHei", size=10.5)
        col_width = pdf.epw / 4
        
        # 表头
        pdf.cell(col_width, 10, "姓名", border=1, align="C")
        pdf.cell(col_width, 10, "性别", border=1, align="C")
        pdf.cell(col_width, 10, "年龄", border=1, align="C")
        pdf.cell(col_width, 10, "病历号", border=1, align="C")
        pdf.ln()
        
        # 数据行
        pdf.set_font("SimSun", size=10.5)
        pdf.cell(col_width, 10, str(case.get("patient_name", "")), border=1, align="C")
        pdf.cell(col_width, 10, str(case.get("gender", "")), border=1, align="C")
        pdf.cell(col_width, 10, f"{case.get('age', '')}岁", border=1, align="C")
        pdf.cell(col_width, 10, str(case.get("case_id", "")), border=1, align="C")
        pdf.ln(10)
        
        # 就诊日期
        visit_date = case.get('visit_date', datetime.now().strftime('%Y-%m-%d'))
        pdf.set_font("SimSun", size=10.5)
        pdf.cell(0, 10, f"就诊日期：{visit_date}", ln=True)
        
        # 分隔线
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + pdf.epw, pdf.get_y())
        pdf.ln(5)
        
        # 各个章节
        if "markdown_content" in case and case["markdown_content"]:
            pdf.set_font("SimSun", size=10.5)
            clean_content = self._clean_markdown(case["markdown_content"])
            pdf.multi_cell(0, 6, clean_content)
            pdf.ln(5)
            
            if "ai_suggestions" in case and case["ai_suggestions"]:
                pdf.set_font("SimHei", size=11)
                pdf.write(5, "【AI 临床建议】")
                pdf.ln(2)
                pdf.set_font("SimSun", size=10.5)
                pdf.set_text_color(22, 160, 133) # 墨绿色
                clean_suggestions = self._clean_markdown(case["ai_suggestions"])
                pdf.multi_cell(0, 6, clean_suggestions)
                pdf.set_text_color(0, 0, 0)
                pdf.ln(5)
        else:
            sections = [
                ("主诉", case.get("chief_complaint", "")),
                ("现病史", case.get("present_illness", "")),
                ("既往史", case.get("past_history", "")),
                ("过敏史", case.get("allergies", "")),
                ("体格检查", case.get("physical_exam", "")),
                ("诊断", case.get("diagnosis", "")),
                ("治疗方案", case.get("treatment_plan", ""))
            ]
            
            for title, content in sections:
                if content:
                    pdf.set_font("SimHei", size=11)
                    pdf.write(5, f"【{title}】")
                    pdf.set_font("SimSun", size=10.5)
                    # 清理可能存在的 Markdown
                    clean_text = self._clean_markdown(content)
                    pdf.multi_cell(0, 6, f" {clean_text}")
                    pdf.ln(3)
        
        # 页脚
        pdf.set_y(-40)
        doctor_name = self.config.get("doctor_name", "王医生")
        pdf.set_font("SimHei", size=11)
        pdf.cell(0, 10, f"医生签名：{doctor_name} __________", ln=True, align="R")
        
        # 优先使用病例保存时的原始时间戳，没有则使用当前时间
        auth_time = case.get("timestamp") or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        pdf.set_font("SimSun", size=9)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 10, f"存证时间戳：{auth_time}", ln=True, align="R")
        
        pdf.output(filepath)
        return filepath

    def _setup_document_styles(self, doc):
        # 设置正文字体
        style = doc.styles['Normal']
        style.font.name = 'SimSun'
        style.font.size = Pt(10.5)  # 五号字
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 自定义标题样式
        if 'Heading 1' not in doc.styles:
            doc.styles.add_style('Heading 1', 1)
        h1 = doc.styles['Heading 1']
        h1.font.name = 'SimHei'
        h1.font.bold = True
        h1.font.size = Pt(14)
        h1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def _add_header(self, doc):
        hospital_name = self.config.get("hospital_name", "XX社区卫生服务中心")
        
        # 医院名称
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(hospital_name)
        run1.font.size = Pt(18)
        run1.font.bold = True
        run1.font.name = 'SimHei'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 病历标题
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("门诊病历")
        run2.font.size = Pt(16)
        run2.font.bold = True
        run2.font.name = 'SimHei'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        doc.add_paragraph()

    def _add_patient_info(self, doc, case):
        # 创建一个更紧凑的表格
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        
        # 设置列宽
        widths = [Inches(1), Inches(1.5), Inches(1), Inches(2)]
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width

        labels = ["姓名", "性别", "年龄", "病历号"]
        values = [
            case.get("patient_name", ""),
            case.get("gender", ""),
            f"{case.get('age', '')}岁",
            case.get("case_id", "")
        ]
        
        for i in range(4):
            # 第一行：标签
            cell_label = table.rows[0].cells[i]
            cell_label.text = labels[i]
            para_label = cell_label.paragraphs[0]
            para_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_label = para_label.runs[0]
            run_label.font.bold = True
            
            # 第二行：数值
            cell_val = table.rows[1].cells[i]
            cell_val.text = values[i]
            para_val = cell_val.paragraphs[0]
            para_val.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加就诊时间行
        doc.add_paragraph(f"就诊日期：{case.get('visit_date', datetime.now().strftime('%Y-%m-%d'))}")

    def _add_divider(self, doc):
        p = doc.add_paragraph()
        run = p.add_run("-" * 80)
        run.font.color.rgb = RGBColor(128, 128, 128)

    def _add_case_content(self, doc, case):
        # 优先使用 markdown_content
        if "markdown_content" in case and case["markdown_content"]:
            # 清理 Markdown 符号
            clean_content = self._clean_markdown(case["markdown_content"])
            p = doc.add_paragraph()
            run = p.add_run(clean_content)
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 如果有 AI 建议，也加上
            if "ai_suggestions" in case and case["ai_suggestions"]:
                doc.add_paragraph()
                p = doc.add_paragraph()
                run_title = p.add_run("【AI 临床建议】")
                run_title.font.bold = True
                run_title.font.name = 'SimHei'
                run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                
                p = doc.add_paragraph()
                # 同样清理 AI 建议中的 Markdown
                clean_suggestions = self._clean_markdown(case["ai_suggestions"])
                run_sug = p.add_run(clean_suggestions)
                run_sug.font.name = 'SimSun'
                run_sug._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run_sug.font.italic = True
            return

        # 兼容旧版本
        sections = [
            ("主诉", case.get("chief_complaint", "")),
            ("现病史", case.get("present_illness", "")),
            ("既往史", case.get("past_history", "")),
            ("过敏史", case.get("allergies", "")),
            ("体格检查", case.get("physical_exam", "")),
            ("诊断", case.get("diagnosis", "")),
            ("治疗方案", case.get("treatment_plan", ""))
        ]
        
        for title, content in sections:
            if content:
                # 标题部分
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                run_title = p.add_run(f"【{title}】")
                run_title.font.bold = True
                run_title.font.name = 'SimHei'
                run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                
                # 内容部分
                # 清理可能存在的 Markdown
                clean_content = self._clean_markdown(content)
                content_lines = clean_content.split('\n')
                for i, line in enumerate(content_lines):
                    if i == 0:
                        run_content = p.add_run(f" {line}")
                    else:
                        p = doc.add_paragraph()
                        run_content = p.add_run(line)
                    
                    run_content.font.name = 'SimSun'
                    run_content._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _add_footer(self, doc, case):
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 签名行
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doctor_name = self.config.get("doctor_name", "王医生")
        run = p.add_run(f"医生签名：{doctor_name} __________")
        run.font.bold = True
        
        # 存证时间戳
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        auth_time = case.get("timestamp") or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        run_date = p_date.add_run(f"存证时间戳：{auth_time}")
        run_date.font.size = Pt(9)
        run_date.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()
