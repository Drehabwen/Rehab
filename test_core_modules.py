import unittest
import json
import os
import sys
from unittest.mock import Mock, patch, MagicMock

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from case_structurer import CaseStructurer
from case_manager import CaseManager
from document_generator import DocumentGenerator


class TestCaseStructurer(unittest.TestCase):
    def setUp(self):
        self.mock_nlp = MagicMock()
        self.structurer = CaseStructurer(self.mock_nlp)

    def test_analyze_dialogue_success(self):
        # 模拟模型响应
        self.mock_nlp.model_base.chat.return_value = {
            "success": True,
            "content": '[{"speaker": "医生", "text": "你好"}, {"speaker": "患者", "text": "头痛"}]'
        }
        
        result = self.structurer.analyze_dialogue("医生说你好患者说头痛")
        
        self.assertEqual(len(result), 2)
        self.assertEqual(result[0]["speaker"], "医生")
        self.assertEqual(result[1]["text"], "头痛")

    def test_structure_success(self):
        # 模拟模型响应
        expected_json = {
            "patient_name": "张三",
            "gender": "男",
            "age": "45",
            "markdown_content": "# 病例正文\n主诉：头痛",
            "ai_suggestions": "建议检查头颅CT"
        }
        self.mock_nlp.model_pro.chat.return_value = {
            "success": True,
            "content": json.dumps(expected_json)
        }
        
        dialogues = [{"speaker": "医生", "text": "你好"}, {"speaker": "患者", "text": "头痛"}]
        result = self.structurer.structure(dialogues)
        
        self.assertEqual(result["patient_name"], "张三")
        self.assertEqual(result["gender"], "男")
        self.assertIn("病例正文", result["markdown_content"])
        self.assertEqual(result["ai_suggestions"], "建议检查头颅CT")

    def test_generate_report_success(self):
        self.mock_nlp.model_pro.chat.return_value = {
            "success": True,
            "content": "正式病历报告内容"
        }
        
        case_data = {"patient_name": "张三"}
        config = {"hospital_name": "测试医院", "doctor_name": "王医生"}
        result = self.structurer.generate_report(case_data, config)
        
        self.assertEqual(result, "正式病历报告内容")


class TestCaseManager(unittest.TestCase):
    def setUp(self):
        self.test_config = {
            "hospital_name": "测试医院",
            "doctor_name": "测试医生",
            "cases_dir": "./test_cases",
            "exports_dir": "./test_exports"
        }
        self.case_manager = CaseManager(self.test_config)
        
        self.test_case = {
            "case_id": "20260125_001",
            "patient_name": "张三",
            "gender": "男",
            "age": 45,
            "visit_date": "2026-01-25",
            "markdown_content": "# 病例正文\n主诉：头痛3天",
            "ai_suggestions": "AI 建议内容",
            "diagnosis": "高血压病"
        }

    def tearDown(self):
        import shutil
        if os.path.exists("./test_cases"):
            shutil.rmtree("./test_cases")
        if os.path.exists("./test_exports"):
            shutil.rmtree("./test_exports")

    def test_create_new_case(self):
        new_case = self.case_manager.create_new_case("李四", "女", 30)
        
        self.assertEqual(new_case["patient_name"], "李四")
        self.assertEqual(new_case["gender"], "女")
        self.assertEqual(new_case["age"], 30)
        self.assertIsNotNone(new_case["case_id"])
        self.assertIsNotNone(new_case["visit_date"])

    def test_save_case(self):
        success, result = self.case_manager.save_case(self.test_case)
        
        self.assertTrue(success)
        self.assertIsInstance(result, str)
        
        case_file = os.path.join("./test_cases", f"{self.test_case['case_id']}.json")
        self.assertTrue(os.path.exists(case_file))

    def test_load_case(self):
        self.case_manager.save_case(self.test_case)
        
        loaded_case = self.case_manager.load_case(self.test_case["case_id"])
        
        self.assertEqual(loaded_case["patient_name"], self.test_case["patient_name"])
        self.assertEqual(loaded_case["diagnosis"], self.test_case["diagnosis"])

    def test_delete_case(self):
        self.case_manager.save_case(self.test_case)
        
        success = self.case_manager.delete_case(self.test_case["case_id"])
        
        self.assertTrue(success)
        
        case_file = os.path.join("./test_cases", f"{self.test_case['case_id']}.json")
        self.assertFalse(os.path.exists(case_file))

    def test_list_cases(self):
        self.case_manager.save_case(self.test_case)
        
        cases = self.case_manager.list_cases()
        
        self.assertGreater(len(cases), 0)
        self.assertEqual(cases[0]["case_id"], self.test_case["case_id"])

    def test_search_cases(self):
        self.case_manager.save_case(self.test_case)
        
        cases = self.case_manager.list_cases()
        
        self.assertGreater(len(cases), 0)
        self.assertEqual(cases[0]["patient_name"], "张三")

    def test_validate_case_success(self):
        is_valid, message = self.case_manager._validate_case(self.test_case)
        
        self.assertTrue(is_valid)
        self.assertEqual(message, "")

    def test_validate_case_missing_field(self):
        invalid_case = self.test_case.copy()
        invalid_case["patient_name"] = ""
        
        is_valid, message = self.case_manager._validate_case(invalid_case)
        
        self.assertFalse(is_valid)
        self.assertIn("patient_name", message)


class TestDocumentGenerator(unittest.TestCase):
    def setUp(self):
        self.test_config = {
            "hospital_name": "测试医院",
            "doctor_name": "测试医生",
            "cases_dir": "./test_cases",
            "exports_dir": "./test_exports"
        }
        self.generator = DocumentGenerator(self.test_config)
        
        self.test_case = {
            "case_id": "20260125_001",
            "patient_name": "张三",
            "gender": "男",
            "age": 45,
            "visit_date": "2026-01-25",
            "markdown_content": "# 病例正文\n主诉：头痛3天",
            "ai_suggestions": "AI 建议内容",
            "diagnosis": "高血压病"
        }

    def tearDown(self):
        import shutil
        if os.path.exists("./test_cases"):
            shutil.rmtree("./test_cases")
        if os.path.exists("./test_exports"):
            shutil.rmtree("./test_exports")

    def test_generate_word(self):
        filepath = self.generator.generate_word(self.test_case)
        
        self.assertTrue(os.path.exists(filepath))
        self.assertIn("20260125_001", filepath)
        self.assertIn("张三", filepath)

    def test_generate_word_content(self):
        filepath = self.generator.generate_word(self.test_case)
        
        from docx import Document
        doc = Document(filepath)
        
        # 提取正文文本
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        
        self.assertIn("张三", text)
        self.assertIn("头痛3天", text)
        self.assertIn("测试医院", text)

    def test_setup_document_styles(self):
        from docx import Document
        doc = Document()
        
        self.generator._setup_document_styles(doc)
        
        self.assertIsNotNone(doc.styles['Normal'].font.name)
        self.assertEqual(doc.styles['Normal'].font.size.pt, 10.5)

    def test_add_header(self):
        from docx import Document
        doc = Document()
        
        self.generator._add_header(doc)
        
        text = doc.paragraphs[0].text
        self.assertIn("测试医院", text)


class TestIntegration(unittest.TestCase):
    def setUp(self):
        self.test_config = {
            "hospital_name": "测试医院",
            "doctor_name": "测试医生",
            "cases_dir": "./test_cases",
            "exports_dir": "./test_exports"
        }
        self.case_manager = CaseManager(self.test_config)
        self.mock_nlp = MagicMock()
        self.structurer = CaseStructurer(self.mock_nlp)
        self.generator = DocumentGenerator(self.test_config)

    def tearDown(self):
        import shutil
        if os.path.exists("./test_cases"):
            shutil.rmtree("./test_cases")
        if os.path.exists("./test_exports"):
            shutil.rmtree("./test_exports")

    def test_full_workflow(self):
        # 1. 模拟 AI 结构化响应
        expected_json = {
            "patient_name": "张三",
            "gender": "男",
            "age": "45",
            "markdown_content": "# 病例正文\n主诉：头痛3天",
            "ai_suggestions": "AI 建议内容"
        }
        self.mock_nlp.model_pro.chat.return_value = {
            "success": True,
            "content": json.dumps(expected_json)
        }
        
        dialogues = [{"speaker": "医生", "text": "你好"}, {"speaker": "患者", "text": "头痛"}]
        structured_case = self.structurer.structure(dialogues)
        
        self.assertEqual(structured_case["patient_name"], "张三")
        
        # 2. 保存病例
        # 添加一些必要字段
        structured_case["case_id"] = "TEST_ID"
        structured_case["diagnosis"] = "高血压"
        success, result = self.case_manager.save_case(structured_case)
        self.assertTrue(success)
        
        # 3. 导出 Word
        filepath = self.generator.generate_word(structured_case)
        self.assertTrue(os.path.exists(filepath))
        
        from docx import Document
        doc = Document(filepath)
        text = ""
        for p in doc.paragraphs: text += p.text + "\n"
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells: text += c.text + "\n"
        
        self.assertIn("张三", text)
        self.assertIn("测试医院", text)

    def test_multiple_cases_workflow(self):
        cases = [
            {
                "patient_name": "张三", "gender": "男", "age": 45,
                "markdown_content": "头痛3天", "diagnosis": "高血压病",
                "ai_suggestions": "建议检查"
            },
            {
                "patient_name": "李四", "gender": "女", "age": 30,
                "markdown_content": "发热1天", "diagnosis": "上呼吸道感染",
                "ai_suggestions": "建议休息"
            }
        ]
        
        saved_cases = []
        for case_data in cases:
            new_case = self.case_manager.create_new_case(
                case_data["patient_name"],
                case_data["gender"],
                case_data["age"]
            )
            new_case.update(case_data)
            success, _ = self.case_manager.save_case(new_case)
            if success:
                saved_cases.append(new_case)
        
        self.assertEqual(len(saved_cases), 2)
        
        all_cases = self.case_manager.list_cases()
        self.assertEqual(len(all_cases), 2)
        
        for case in saved_cases:
            filepath = self.generator.generate_word(case)
            self.assertTrue(os.path.exists(filepath))


def run_tests():
    loader = unittest.TestLoader()
    suite = unittest.TestSuite()
    
    suite.addTests(loader.loadTestsFromTestCase(TestCaseStructurer))
    suite.addTests(loader.loadTestsFromTestCase(TestCaseManager))
    suite.addTests(loader.loadTestsFromTestCase(TestDocumentGenerator))
    suite.addTests(loader.loadTestsFromTestCase(TestIntegration))
    
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    
    return result.wasSuccessful()


if __name__ == "__main__":
    success = run_tests()
    sys.exit(0 if success else 1)
