#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
语音转病例助手 - 快速测试脚本
用于快速验证核心功能
"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))


def test_case_structurer():
    """测试病例结构化功能"""
    print("="*60)
    print("测试 1: 病例结构化功能")
    print("="*60)
    
    from case_structurer import CaseStructurer
    
    structurer = CaseStructurer()
    
    transcript = """
    主诉：头痛3天
    现病史：患者3天前无明显诱因出现头痛
    既往史：高血压病史5年
    过敏史：青霉素过敏
    体格检查：T 36.5℃，BP 140/90mmHg
    诊断：高血压病
    治疗：继续降压治疗
    """
    
    result = structurer.structure(transcript)
    
    print(f"✓ 主诉: {result['chief_complaint']}")
    print(f"✓ 现病史: {result['present_illness'][:30]}...")
    print(f"✓ 既往史: {result['past_history']}")
    print(f"✓ 过敏史: {result['allergies']}")
    print(f"✓ 体格检查: {result['physical_exam']}")
    print(f"✓ 诊断: {result['diagnosis']}")
    print(f"✓ 治疗: {result['treatment_plan']}")
    
    assert result['chief_complaint'] == "头痛3天", "主诉提取失败"
    assert "高血压" in result['past_history'], "既往史提取失败"
    assert result['diagnosis'] == "高血压病", "诊断提取失败"
    
    print("\n✅ 病例结构化测试通过\n")
    return True


def test_case_manager():
    """测试病例管理功能"""
    print("="*60)
    print("测试 2: 病例管理功能")
    print("="*60)
    
    from case_manager import CaseManager
    
    config = {
        "hospital_name": "测试医院",
        "doctor_name": "测试医生",
        "cases_dir": "./test_cases",
        "exports_dir": "./test_exports"
    }
    
    manager = CaseManager(config)
    
    test_case = {
        "case_id": "20260125_001",
        "patient_name": "张三",
        "gender": "男",
        "age": 45,
        "visit_date": "2026-01-25",
        "chief_complaint": "头痛3天",
        "present_illness": "患者3天前无明显诱因出现头痛",
        "past_history": "高血压病史5年",
        "allergies": "青霉素过敏",
        "physical_exam": "T 36.5℃，BP 140/90mmHg",
        "diagnosis": "高血压病",
        "treatment_plan": "继续降压治疗"
    }
    
    print("✓ 创建新病例...")
    new_case = manager.create_new_case("李四", "女", 30)
    print(f"  病例ID: {new_case['case_id']}")
    print(f"  患者姓名: {new_case['patient_name']}")
    
    print("✓ 保存病例...")
    success, result = manager.save_case(test_case)
    print(f"  保存结果: {result}")
    assert success, "保存病例失败"
    
    print("✓ 加载病例...")
    loaded_case = manager.load_case(test_case['case_id'])
    print(f"  加载的患者: {loaded_case['patient_name']}")
    print(f"  加载的诊断: {loaded_case['diagnosis']}")
    assert loaded_case is not None, "加载病例失败"
    assert loaded_case['patient_name'] == "张三", "患者姓名不匹配"
    
    print("✓ 列出病例...")
    cases = manager.list_cases()
    print(f"  病例数量: {len(cases)}")
    assert len(cases) > 0, "列出病例失败"
    
    print("✓ 删除病例...")
    success = manager.delete_case(test_case['case_id'])
    print(f"  删除结果: {success}")
    assert success, "删除病例失败"
    
    print("\n✅ 病例管理测试通过\n")
    return True


def test_document_generator():
    """测试文档生成功能"""
    print("="*60)
    print("测试 3: 文档生成功能")
    print("="*60)
    
    from document_generator import DocumentGenerator
    
    config = {
        "hospital_name": "测试医院",
        "doctor_name": "测试医生",
        "cases_dir": "./test_cases",
        "exports_dir": "./test_exports"
    }
    
    generator = DocumentGenerator(config)
    
    test_case = {
        "case_id": "20260125_001",
        "patient_name": "张三",
        "gender": "男",
        "age": 45,
        "visit_date": "2026-01-25",
        "chief_complaint": "头痛3天",
        "present_illness": "患者3天前无明显诱因出现头痛",
        "past_history": "高血压病史5年",
        "allergies": "青霉素过敏",
        "physical_exam": "T 36.5℃，BP 140/90mmHg",
        "diagnosis": "高血压病",
        "treatment_plan": "继续降压治疗"
    }
    
    print("✓ 生成Word文档...")
    filepath = generator.generate_word(test_case)
    print(f"  文件路径: {filepath}")
    
    assert os.path.exists(filepath), "Word文档未生成"
    
    print("✓ 验证文档内容...")
    from docx import Document
    doc = Document(filepath)
    
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    
    assert "张三" in text, "文档中缺少患者姓名"
    assert "头痛3天" in text, "文档中缺少主诉"
    assert "高血压病" in text, "文档中缺少诊断"
    assert "测试医院" in text, "文档中缺少医院名称"
    
    print("  ✓ 包含患者姓名: 张三")
    print("  ✓ 包含主诉: 头痛3天")
    print("  ✓ 包含诊断: 高血压病")
    print("  ✓ 包含医院名称: 测试医院")
    
    print("\n✅ 文档生成测试通过\n")
    return True


def test_integration():
    """测试完整工作流程"""
    print("="*60)
    print("测试 4: 完整工作流程")
    print("="*60)
    
    from case_structurer import CaseStructurer
    from case_manager import CaseManager
    from document_generator import DocumentGenerator
    
    config = {
        "hospital_name": "测试医院",
        "doctor_name": "测试医生",
        "cases_dir": "./test_cases",
        "exports_dir": "./test_exports"
    }
    
    structurer = CaseStructurer()
    manager = CaseManager(config)
    generator = DocumentGenerator(config)
    
    transcript = """
    主诉：头痛3天
    现病史：患者3天前无明显诱因出现头痛，伴恶心
    既往史：高血压病史5年
    过敏史：青霉素过敏
    体格检查：T 36.5℃，BP 140/90mmHg
    诊断：高血压病
    治疗：继续降压治疗，监测血压
    """
    
    print("✓ 步骤1: 结构化病例...")
    structured_case = structurer.structure(transcript)
    print(f"  主诉: {structured_case['chief_complaint']}")
    print(f"  诊断: {structured_case['diagnosis']}")
    
    print("✓ 步骤2: 保存病例...")
    structured_case['case_id'] = manager._generate_case_id()
    structured_case['patient_name'] = "张三"
    structured_case['gender'] = "男"
    structured_case['age'] = 45
    structured_case['visit_date'] = "2026-01-25"
    success, result = manager.save_case(structured_case)
    print(f"  保存结果: {result}")
    assert success, "保存病例失败"
    
    print("✓ 步骤3: 生成Word文档...")
    filepath = generator.generate_word(structured_case)
    print(f"  文件路径: {filepath}")
    assert os.path.exists(filepath), "Word文档未生成"
    
    print("✓ 步骤4: 验证文档...")
    from docx import Document
    doc = Document(filepath)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    assert "头痛3天" in text, "文档中缺少主诉"
    assert "高血压病" in text, "文档中缺少诊断"
    assert "测试医院" in text, "文档中缺少医院名称"
    
    print("  ✓ 文档包含所有关键信息")
    
    print("\n✅ 完整工作流程测试通过\n")
    return True


def cleanup():
    """清理测试文件"""
    print("="*60)
    print("清理测试文件")
    print("="*60)
    
    import shutil
    
    if os.path.exists("./test_cases"):
        shutil.rmtree("./test_cases")
        print("✓ 删除 test_cases 目录")
    
    if os.path.exists("./test_exports"):
        shutil.rmtree("./test_exports")
        print("✓ 删除 test_exports 目录")
    
    print()


def main():
    """主函数"""
    print("\n" + "="*60)
    print("语音转病例助手 - 快速测试")
    print("="*60 + "\n")
    
    tests = [
        ("病例结构化", test_case_structurer),
        ("病例管理", test_case_manager),
        ("文档生成", test_document_generator),
        ("完整工作流程", test_integration),
    ]
    
    results = []
    for test_name, test_func in tests:
        try:
            success = test_func()
            results.append((test_name, success, None))
        except Exception as e:
            print(f"❌ {test_name} 测试失败: {e}\n")
            results.append((test_name, False, str(e)))
    
    cleanup()
    
    print("="*60)
    print("测试结果汇总")
    print("="*60)
    
    passed = sum(1 for _, success, _ in results if success)
    failed = len(results) - passed
    
    for test_name, success, error in results:
        status = "✅ 通过" if success else "❌ 失败"
        print(f"{status} - {test_name}")
        if error:
            print(f"     错误: {error}")
    
    print()
    print(f"总计: {len(results)} 个测试")
    print(f"通过: {passed} 个")
    print(f"失败: {failed} 个")
    print("="*60)
    
    if failed == 0:
        print("\n🎉 所有测试通过！\n")
        return 0
    else:
        print(f"\n⚠️  有 {failed} 个测试失败\n")
        return 1


if __name__ == "__main__":
    sys.exit(main())
